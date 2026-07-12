"""
考研伴学 Agent - 文档解析与切分工具
====================================
支持的格式：PDF / Markdown / TXT
切分策略：优先按 Markdown 标题层级切分（保留知识结构），
        文本层级不够时回退为字符递归切分。
"""

from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)

from app.core.config import settings


# ============================================
# 第〇部分：前置内容过滤
# ============================================

# 前置内容关键词（页面前几段如果匹配这些词，整页丢弃）
FRONT_MATTER_PATTERNS = [
    "前言", "序言", "致读者", "出版说明", "编者的话",
    "使用说明", "导读", "写在前面", "作者简介",
    "内容提要", "图书在版编目", "CIP数据", "版权信息",
]


def _looks_like_front_matter(text: str) -> bool:
    """
    判断文本是否像前置内容（扉页/前言/版权页等）。

    启发式规则:
      1. 文本开头 200 字符内匹配 FRONT_MATTER_PATTERNS 关键词
      2. 文本极短（< 30 字符），通常为章节标题页
    """
    if not text or not text.strip():
        return True
    if len(text.strip()) < 30:
        return True
    head = text[:200]
    return any(p in head for p in FRONT_MATTER_PATTERNS)


# ============================================
# 第一部分：文档加载
# ============================================

def load_document(file_path: str) -> List[Document]:
    """
    根据文件后缀自动选择合适的加载器，将文件加载为 LangChain Document 列表。

    参数:
        file_path: 文件路径（支持 .pdf / .markdown / .txt / .docx）

    返回:
        List[Document] —— 每个 Document 的 page_content 为提取后的文本

    异常:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        loader = PyPDFLoader(str(path))
        # [缺陷] PyPDFLoader 只能提取文本层 PDF，无法处理扫描版（需要 OCR）。
        # [缺陷] 不提取表格、公式、图片，考研数学/理工类资料会丢失大量信息。
        # [后续扩展] 可替换为 Marker / Nougat / Mathpix 等专业工具。
        docs = loader.load()
        # 为每个 page 补上来源元数据
        for i, doc in enumerate(docs):
            doc.metadata["source"] = str(path)
            doc.metadata["page"] = i + 1
        return docs

    elif suffix == ".markdown":
        loader = TextLoader(str(path), encoding="utf-8")
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = str(path)
        return docs

    elif suffix == ".txt":
        loader = TextLoader(str(path), encoding="utf-8")
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = str(path)
        return docs

    elif suffix == ".docx":
        try:
            import docx
            doc = docx.Document(str(path))
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if not full_text.strip():
                raise RuntimeError("文档为空或仅含图片")
            docs = [
                Document(
                    page_content=full_text,
                    metadata={"source": str(path)},
                )
            ]
            return docs
        except ImportError:
            raise ImportError(
                "处理 .docx 文件需要安装 python-docx，请运行: pip install python-docx"
            )


    else:
        raise ValueError(
            f"不支持的文件格式: '{suffix}'。"
            f"当前支持: .pdf, .markdown, .txt, .docx"
        )


# ============================================
# 第二部分：Markdown 标题层级切分
# ============================================

# 根据配置动态生成 LangChain 支持的标题分隔符
# MARKDOWN_HEADER_MAX_LEVEL=4 → ["#", "##", "###", "####"]
_HEADERS_TO_SPLIT_ON = [
    ("#" * i, f"Header_{i}")
    for i in range(1, settings.MARKDOWN_HEADER_MAX_LEVEL + 1)
]


def _build_header_splitter() -> MarkdownHeaderTextSplitter:
    """
    构建 Markdown 标题切分器。
    每次调用都从 settings 读取最新配置，保证灵活性。
    """
    # 只保留切分层级以下的标题（如 SPLIT_LEVEL=2，则只按 H1、H2 切分）
    effective_headers = _HEADERS_TO_SPLIT_ON[:settings.MARKDOWN_SPLIT_HEADER_LEVEL]

    return MarkdownHeaderTextSplitter(
        headers_to_split_on=effective_headers,
        # strip_headers=False 表示保留标题行原文（不删除）
        strip_headers=False,
    )


def split_by_markdown_headers(docs: List[Document]) -> List[Document]:
    """
    按 Markdown 标题层级切分文档。
    适用于已经有标题结构的教材笔记（.markdown 或从 PDF 转 Markdown 后的内容）。

    切分策略:
        1. 先用 MarkdownHeaderTextSplitter 按标题边界切分（保证不跨章节）。
        2. 如果某块仍超过 CHUNK_SIZE，再用 RecursiveCharacterTextSplitter 做二级切分。

    参数:
        docs: 待切分的 Document 列表

    返回:
        切分并打上元数据标签后的 Document 列表
    """
    header_splitter = _build_header_splitter()
    # 二级切分器 —— 当某个块仍然过长时使用
    chunk_splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
                separators=[
            "\n\n",     # 空行（段落边界）
            "\n",       # 换行
            "$$",       # 独立公式块结束（LaTeX 数学环境边界）
            "。",       # 中文句号
            ". ",       # 英文句号后空格
            "；",       # 中文分号
            "; ",       # 英文分号后空格
            "，",       # 中文逗号
            ", ",       # 英文逗号后空格
            " ",        # 最终兜底：按空格拆
        ],

        # [缺陷] 当前分隔符仅适配中英文混合，未针对数学公式、代码块优化。
    )

    final_docs: List[Document] = []

    for doc in docs:
        # 步骤1: Markdown 标题切分（最优策略）
        try:
            header_chunks = header_splitter.split_text(doc.page_content)
        except Exception:
            header_chunks = []

        # 步骤2: 标题切分未生效 → 退化为段落切分（中间层兜底）
        # MarkdownHeaderTextSplitter 即使没标题也返回 1 个 Document，
        # 所以需要检查切出来的块是否真的携带了标题元数据
        has_header_metadata = any(
            chunk_doc.metadata.get("Header_1") for chunk_doc in header_chunks
        )
        if not has_header_metadata:
            paragraphs = [p.strip() for p in doc.page_content.split("\n\n") if p.strip()]

            if paragraphs:
                header_chunks = [
                    Document(
                        page_content=p,
                        metadata={**doc.metadata, "split_method": "paragraph"},
                    )
                    for p in paragraphs
                ]
            else:
                # 步骤3: 段落也分不出来 → 最终退化：整个文档作为一个块
                header_chunks = [
                    Document(
                        page_content=doc.page_content,
                        metadata={**doc.metadata, "split_method": "raw"},
                    )
                ]
        # 步骤4: 对每个块做长度检查，超长再切
        for chunk_doc in header_chunks:
            chunk_text = chunk_doc.page_content.strip()
            if not chunk_text:
                continue

            # 合并元数据：原始文档元数据 + 标题切分产生的层级元数据
            # 保留先前设置的切分方法（如 paragraph/raw），未设置时默认 header
            chunk_doc.metadata.setdefault("split_method", "header")
            merged_metadata = {**doc.metadata, **chunk_doc.metadata}

            if len(chunk_text) <= settings.CHUNK_SIZE:
                final_docs.append(
                    Document(
                        page_content=chunk_text,
                        metadata=merged_metadata,
                    )
                )
            else:
                # 超长块：递归字符切分
                sub_chunks = chunk_splitter.split_text(chunk_text)
                for sub in sub_chunks:
                    sub = sub.strip()
                    if not sub:
                        continue
                    final_docs.append(
                        Document(
                            page_content=sub,
                            metadata={
                                **merged_metadata,
                                "split_method": "recursive",
                            },
                        )
                    )

    return final_docs


def process_document(file_path: str) -> List[Document]:
    """
    一键流程: 加载 → 过滤前置内容 → 切分。
    这是文档处理的主要入口，外部只需调用此函数。

    参数:
        file_path: 文件路径

    返回:
        已切分的 Document 列表，可直接送入向量库

    异常:
        RuntimeError: 文件为空或未提取到有效内容
    """
    docs = load_document(file_path)

    # 空文件检测
    if not docs:
        raise RuntimeError(f"文档解析结果为空: {file_path}，请检查文件是否损坏或为空")
    all_empty = all(not d.page_content.strip() for d in docs)
    if all_empty:
        raise RuntimeError(f"文档 '{Path(file_path).name}' 未提取到有效文本内容，可能为扫描版PDF或空文件")

    # 过滤前置内容（扉页/前言/目录页）
    filtered_docs = [d for d in docs if not _looks_like_front_matter(d.page_content)]
    if filtered_docs:
        skipped = len(docs) - len(filtered_docs)
        if skipped > 0:
            print(f"  [过滤] 已跳过 {skipped} 页前置内容（前言/扉页等）")
        docs = filtered_docs

    return split_by_markdown_headers(docs)


